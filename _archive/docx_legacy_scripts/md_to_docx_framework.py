#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Convert AI协作项目全生命周期框架.md to a polished Word document.

Upgraded from the V2.4 converter — adds code blocks, inline code, horizontal rules,
and improved table handling for the framework's complex formatting.
"""

import re
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers ──────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell borders (thin grey)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:color'), kwargs.get('color', 'AAAAAA'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_run_font(run, font_name='微软雅黑', size=10.5, bold=False, italic=False,
                 color=None, mono=False):
    font = run.font
    if mono:
        font.name = 'Consolas'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    else:
        font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    font.size = Pt(size)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = RGBColor(*color)

def add_heading_custom(doc, text, level):
    p = doc.add_heading(level=level)
    run = p.add_run(text)
    sizes = {1: 18, 2: 15, 3: 13, 4: 11.5}
    set_run_font(run, size=sizes.get(level, 11.5), bold=True,
                 color=(0x1A, 0x23, 0x7E))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_para(doc, text='', bold=False, indent=False, mono=False, size=10.5,
             italic=False, color=None, space_after=3, space_before=3,
             alignment=None):
    """Add a paragraph with inline **bold** and `code` parsing."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if alignment is not None:
        p.alignment = alignment

    if not text:
        return p

    # Split by inline code first, then bold within each segment
    segments = re.split(r'(`[^`]+`)', text)
    for seg in segments:
        if seg.startswith('`') and seg.endswith('`'):
            run = p.add_run(seg[1:-1])
            set_run_font(run, mono=True, size=size - 0.5)
            continue
        # Bold within this segment
        parts = re.split(r'\*\*(.*?)\*\*', seg)
        for idx, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            is_bold = bold or (idx % 2 == 1)
            set_run_font(run, bold=is_bold, italic=italic, color=color,
                         size=size)
    return p

def add_code_block(doc, code_lines):
    """Add a monospace code block with light grey background."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        # Light grey background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'F5F5F5')
        shd.set(qn('w:val'), 'clear')
        pPr.append(shd)
        run = p.add_run(line if line else ' ')
        set_run_font(run, mono=True, size=8.5)
    # Small spacer after code block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(2)
    spacer.paragraph_format.space_after = Pt(2)

def parse_table(lines):
    """Parse markdown table lines into rows."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        # Skip separator rows like |---|:---:|---|
        if all(set(c) <= set(' :-') for c in cells):
            continue
        rows.append(cells)
    return rows

def add_table_custom(doc, rows):
    """Add a styled table. Header row gets blue background."""
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(rows):
        for j in range(num_cols):
            cell = table.rows[i].cells[j]
            cell.text = ''
            if j < len(row):
                cell.text = row[j]
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(1)
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    set_run_font(run, size=8.5)
            if i == 0:
                set_cell_shading(cell, 'D5E8F0')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_run_font(run, size=8.5, bold=True)
    doc.add_paragraph()  # spacer after table


def insert_toc_field(doc):
    """Insert a native Word TOC field that reads Heading 1-3 styles.
    User must right-click → Update Field (or press F9) in Word to populate page numbers."""
    toc_paragraph = doc.add_paragraph()
    toc_paragraph.style = doc.styles['Normal']
    run = toc_paragraph.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)
    run2 = toc_paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)
    run3 = toc_paragraph.add_run()
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar_separate)
    run4 = toc_paragraph.add_run()
    run4_text = OxmlElement('w:t')
    run4_text.text = '（在 Word 中右键此处 → 更新域，即可自动生成带页码的目录）'
    run4._r.append(run4_text)
    run5 = toc_paragraph.add_run()
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar_end)
    doc.add_paragraph()  # spacer

def process_markdown(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    doc = Document()

    # A4 page
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    lines = content.splitlines()
    i = 0
    table_lines = []
    in_table = False
    code_lines = []
    in_code_block = False
    in_toc = False  # skip the manual markdown TOC (Word auto-generates its own)

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # ── skip manual TOC (Word auto-generates its own with page numbers) ─
        if in_toc:
            if stripped == '---':
                in_toc = False
            i += 1
            continue
        if stripped == '## 目录':
            insert_toc_field(doc)
            in_toc = True
            i += 1
            continue

        # ── code block ──────────────────────────────────────────
        if stripped.startswith('```'):
            if in_code_block:
                # End code block
                if code_lines:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── finish pending table ────────────────────────────────
        if in_table and not stripped.startswith('|'):
            in_table = False
            rows = parse_table(table_lines)
            add_table_custom(doc, rows)
            table_lines = []

        # ── horizontal rule ──────────────────────────────────────
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin horizontal line via bottom border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── headings ─────────────────────────────────────────────
        if stripped.startswith('# '):
            add_heading_custom(doc, stripped[2:], 1)
            i += 1
            continue
        elif stripped.startswith('## '):
            add_heading_custom(doc, stripped[3:], 2)
            i += 1
            continue
        elif stripped.startswith('### '):
            add_heading_custom(doc, stripped[4:], 3)
            i += 1
            continue
        elif stripped.startswith('#### '):
            add_heading_custom(doc, stripped[5:], 4)
            i += 1
            continue

        # ── table ────────────────────────────────────────────────
        if stripped.startswith('|'):
            in_table = True
            table_lines.append(stripped)
            i += 1
            continue

        # ── empty line ───────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        # ── blockquote ───────────────────────────────────────────
        if stripped.startswith('> '):
            # Check for bold inside blockquote
            text = stripped[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.74)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                if not part:
                    continue
                run = p.add_run(part)
                set_run_font(run, italic=True, bold=(idx % 2 == 1),
                             color=(0x55, 0x55, 0x55), size=10)
            i += 1
            continue

        # ── list items ───────────────────────────────────────────
        list_match = re.match(r'^(\s*)([-*]|\d+\.)\s+(.*)', stripped)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            text = list_match.group(3)
            style = 'List Bullet' if list_match.group(2) in ('-', '*') else 'List Number'
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent = Cm(0.74 + indent_level * 0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            parts = re.split(r'\*\*(.*?)\*\*', text)
            for idx, part in enumerate(parts):
                if not part:
                    continue
                # Handle inline code within list items
                sub_parts = re.split(r'(`[^`]+`)', part)
                for sp in sub_parts:
                    if sp.startswith('`') and sp.endswith('`'):
                        run = p.add_run(sp[1:-1])
                        set_run_font(run, mono=True, size=9.5)
                    else:
                        run = p.add_run(sp)
                        set_run_font(run, bold=(idx % 2 == 1))
            i += 1
            continue

        # ── normal paragraph ─────────────────────────────────────
        add_para(doc, stripped)
        i += 1

    # Handle trailing table or code block
    if table_lines:
        rows = parse_table(table_lines)
        add_table_custom(doc, rows)
    if code_lines:
        add_code_block(doc, code_lines)

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == '__main__':
    md = r'<PROJECT_ROOT>/AI协作项目全生命周期框架.md'
    docx = r'<PROJECT_ROOT>/AI协作项目全生命周期框架.docx'
    process_markdown(md, docx)
