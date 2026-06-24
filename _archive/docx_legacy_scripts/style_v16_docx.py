#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Post-process pandoc .docx: 微软雅黑 fonts + Table Grid style + Word TOC field."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# If run standalone, use _v16_styled; if called from embed_mermaid_png, use final .docx
import os as _os
DOCX = _os.environ.get('STYLE_DOCX_PATH',
    "<PROJECT_ROOT>/AI协作项目全生命周期框架.docx")

def apply_font(run, name='微软雅黑', size=None, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if color:
        from docx.shared import RGBColor
        run.font.color.rgb = RGBColor(*color)

def add_table_grid(table):
    """Apply Table Grid style and borders to a table."""
    # Set style
    table.style = None  # clear any existing style
    tbl = table._element
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Set Table Grid style
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is None:
        tblStyle = OxmlElement('w:tblStyle')
        tblPr.append(tblStyle)
    tblStyle.set(qn('w:val'), 'TableGrid')

    # Add borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def insert_word_toc(doc):
    """Replace hand-written TOC with a Word TOC field (page numbers + hyperlinks)."""
    # Find the 目录 paragraph
    toc_start = None
    toc_end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == '目录' and 'Heading' in (p.style.name or ''):
            toc_start = i
        elif toc_start is not None and 'Heading' in (p.style.name or '') and i > toc_start:
            toc_end = i
            break
    if toc_start is None:
        print("WARNING: TOC heading not found, skipping TOC field insertion")
        return

    # Remove hand-written TOC paragraphs (between toc_start and toc_end)
    to_remove = []
    for i in range(toc_start, toc_end):
        p = doc.paragraphs[i]
        if 'Heading' not in (p.style.name or ''):
            to_remove.append(p)
    for p in to_remove:
        p._element.getparent().remove(p._element)

    # Find the TOC heading again (index may have shifted)
    toc_heading = None
    for p in doc.paragraphs:
        if p.text.strip() == '目录' and 'Heading' in (p.style.name or ''):
            toc_heading = p
            break

    if toc_heading is None:
        print("WARNING: TOC heading lost during cleanup")
        return

    # Insert Word TOC field after the TOC heading
    toc_para = OxmlElement('w:p')
    toc_run = OxmlElement('w:r')

    # Field begin
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    toc_run.append(fldChar_begin)

    # Field instruction
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    toc_run.append(instrText)

    # Field separator
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    toc_run.append(fldChar_sep)

    # Field content (placeholder text)
    placeholder_run = OxmlElement('w:r')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.set(qn('xml:space'), 'preserve')
    placeholder_text.text = '（右键点击此处 → 更新域，即可生成目录）'
    placeholder_run.append(placeholder_text)
    toc_run.append(placeholder_run)

    # Field end
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    toc_run.append(fldChar_end)

    toc_para.append(toc_run)

    # Insert after TOC heading
    toc_heading._element.addnext(toc_para)

    print("Word TOC field inserted (right-click → Update Field to populate)")

# ====== MAIN ======
doc = Document(DOCX)

# 1. Fonts
for p in doc.paragraphs:
    style_name = p.style.name or ''
    for run in p.runs:
        if 'Heading 1' in style_name:
            apply_font(run, '微软雅黑', Pt(18), color=(0x1A, 0x23, 0x7E))
        elif 'Heading 2' in style_name:
            apply_font(run, '微软雅黑', Pt(14), color=(0x1A, 0x23, 0x7E))
        elif 'Heading 3' in style_name:
            apply_font(run, '微软雅黑', Pt(12), color=(0x1A, 0x23, 0x7E))
        elif 'Heading' in style_name:
            apply_font(run, '微软雅黑', Pt(11), color=(0x1A, 0x23, 0x7E))
        elif 'toc' in style_name.lower():
            apply_font(run, '微软雅黑', Pt(10))
        else:
            apply_font(run, '微软雅黑', Pt(10))

# Table cell fonts
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    apply_font(run, '微软雅黑', Pt(9))

# Heading style fonts
for style in doc.styles:
    if style.name and 'Heading' in style.name:
        style.font.name = '微软雅黑'
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), '微软雅黑')
        rFonts.set(qn('w:hAnsi'), '微软雅黑')
        rFonts.set(qn('w:eastAsia'), '微软雅黑')

normal = doc.styles['Normal']
normal.font.name = '微软雅黑'
normal.font.size = Pt(10)

# 2. Table Grid style
for table in doc.tables:
    add_table_grid(table)
print(f"Table Grid applied to {len(doc.tables)} tables")

# 3. Word TOC field
insert_word_toc(doc)

doc.save(DOCX)
print(f"Done: {DOCX} ({os.path.getsize(DOCX)} bytes)")
