#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI协作项目全生命周期框架 — v1.6.3 DOCX 全量重生成 + 边距调整 (2.54cm→2cm)
1. 备份当前根目录 docx
2. 运行 embed_mermaid_png.py 全量 pandoc 管线
3. 调整页面边距 → 2cm
4. 更新 core properties → v1.6.3
5. 输出到根目录
生成模型: DeepSeek-V4-Pro (via Claude Code CLI)
日期: 2026-06-21
"""

import os
import sys
import shutil
import subprocess
from datetime import datetime

BASE = "../AI协作项目全生命周期框架"
DOCX_ROOT = f"{BASE}/AI协作项目全生命周期框架.docx"
PIPELINE_DOCX = f"{BASE}/_pipeline_output/AI协作项目全生命周期框架.docx"
BACKUP_DIR = f"{BASE}/_backups"
EMBED_SCRIPT = f"{BASE}/_archive/docx_legacy_scripts/embed_mermaid_png.py"

os.chdir(BASE)

def step(msg):
    print(f"\n{'='*60}")
    print(f"  {msg}")
    print(f"{'='*60}")

# ====== 1. Backup ======
step("1. 备份当前根目录 DOCX")
backup_name = f"AI协作项目全生命周期框架.docx.v162_pre_v163_sync.backup"
backup_path = os.path.join(BACKUP_DIR, backup_name)
if os.path.exists(DOCX_ROOT):
    shutil.copy2(DOCX_ROOT, backup_path)
    print(f"  Backup: {backup_path} ({os.path.getsize(DOCX_ROOT):,} bytes)")
else:
    print("  WARNING: Root DOCX not found, skipping backup")

# ====== 2. Run full pipeline ======
step("2. 运行 embed_mermaid_png.py 全量 pandoc 管线")
print(f"  Script: {EMBED_SCRIPT}")
# embed_mermaid_png.py uses pypandoc + mmdc + style post-process
# It writes directly to the root DOCX path

# Temporarily rename root DOCX so pipeline writes fresh
temp_backup = DOCX_ROOT + ".temp_backup"
if os.path.exists(DOCX_ROOT):
    os.rename(DOCX_ROOT, temp_backup)

try:
    result = subprocess.run(
        [sys.executable, EMBED_SCRIPT],
        cwd=BASE,
        capture_output=True,
        text=True,
        timeout=600,  # 10 min timeout for Mermaid rendering + pandoc
        env={**os.environ, "PYTHONIOENCODING": "utf-8"}
    )
    print(result.stdout[-2000:] if len(result.stdout) > 2000 else result.stdout)
    if result.returncode != 0:
        print(f"  STDERR: {result.stderr[-1000:]}")
        raise RuntimeError(f"embed_mermaid_png.py failed with code {result.returncode}")
except Exception as e:
    # Restore original
    if os.path.exists(temp_backup):
        os.rename(temp_backup, DOCX_ROOT)
    raise e

# Clean up temp
if os.path.exists(temp_backup):
    os.remove(temp_backup)

if not os.path.exists(DOCX_ROOT):
    # Check if pipeline wrote to _pipeline_output instead
    if os.path.exists(PIPELINE_DOCX):
        shutil.copy2(PIPELINE_DOCX, DOCX_ROOT)
        print(f"  Copied from _pipeline_output/ to root")
    else:
        raise FileNotFoundError(f"DOCX not found at {DOCX_ROOT} or {PIPELINE_DOCX}")

print(f"  Pipeline output: {DOCX_ROOT} ({os.path.getsize(DOCX_ROOT):,} bytes)")

# ====== 3. Adjust page margins → 2cm ======
step("3. 调整页面边距 2.54cm → 2cm")
from docx import Document
from docx.shared import Cm

doc = Document(DOCX_ROOT)

for i, section in enumerate(doc.sections):
    old = (section.top_margin, section.bottom_margin, section.left_margin, section.right_margin)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    print(f"  Section {i}: {old} → (2.0cm, 2.0cm, 2.0cm, 2.0cm)")

# ====== 4. Update version to v1.6.3 ======
step("4. 更新版本元数据 → v1.6.3")
cp = doc.core_properties
cp.version = "v1.6.3"
cp.revision = 19
cp.modified = datetime(2026, 6, 21)
print(f"  version={cp.version}, revision={cp.revision}, modified={cp.modified}")

# Also update version header in early paragraphs
updated_header = False
for i, para in enumerate(doc.paragraphs):
    if i > 20:
        break
    text = para.text
    if 'v1.6' in text and ('版本' in text or 'AI协作项目全生命周期框架' in text):
        for run in para.runs:
            if 'v1.6.2' in run.text:
                run.text = run.text.replace('v1.6.2', 'v1.6.3')
                updated_header = True
                print(f"  Updated v1.6.2→v1.6.3 in paragraph {i}: '{text[:80]}...'")
                break
            elif 'v1.6.1' in run.text:
                run.text = run.text.replace('v1.6.1', 'v1.6.3')
                updated_header = True
                print(f"  Updated v1.6.1→v1.6.3 in paragraph {i}")
                break
if not updated_header:
    print("  NOTE: Could not find version header to update (may use pandoc-generated heading)")

# Update status footer
status_updated = False
for para in reversed(doc.paragraphs):
    if '本文档状态' in para.text or '三件套' in para.text:
        new_status = (
            "本文档状态: v1.6.3 草案 (2026-06-21)。"
            "v1.6.3新增: §1.8局限#9(作者-读者同构假设)+#10(外部依赖漂移风险)"
            "+§2.6规则退役判定(填补只有'加入'无'毕业/退场'的结构性缺口)"
            "+配套外部依赖登记表(.md+.json双件)+可调参数索引(.md)。"
            "经Codex GPT-5.5魔鬼代言人审查(写前)+Qwen qwen3.7-max完备性检查(写后)双后端独立审查写入。"
            "v1.6.2: §7.7被动观测+§9.11跨层可观测性设计。"
            "v1.6.1: A2 Qwen跨模型复现写回。v1.6: 证据体系升级+维护性增强。"
            "⚠️ 本文档从.md源码通过pandoc+Mermaid PNG管线生成。以.md为权威版本。"
            "三件套: .md✅ v1.6.3 / .json✅ v1.6.3 / .docx✅ v1.6.3 (边距2cm,本次重生成)"
        )
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_status
        else:
            para.add_run(new_status)
        status_updated = True
        print(f"  Updated status footer")
        break

if not status_updated:
    # Append to last paragraph
    last_para = doc.paragraphs[-1]
    new_status = (
        "本文档状态: v1.6.3 草案 (2026-06-21)。"
        "v1.6.3新增: §1.8局限#9/#10+§2.6规则退役判定+配套外部依赖登记表+可调参数索引。"
        "经Codex+Qwen双后端独立审查写入。"
        "⚠️ 以.md为权威版本。三件套: .md✅ v1.6.3 / .json✅ v1.6.3 / .docx✅ v1.6.3 (边距2cm)"
    )
    for run in last_para.runs:
        run.text = ""
    if last_para.runs:
        last_para.runs[0].text = new_status
    else:
        last_para.add_run(new_status)
    print(f"  Appended status footer to last paragraph")

# ====== 5. Save ======
step("5. 保存")
doc.save(DOCX_ROOT)
final_size = os.path.getsize(DOCX_ROOT)
print(f"  Saved: {DOCX_ROOT}")
print(f"  Size: {final_size:,} bytes ({final_size/1024:.1f} KB)")

# ====== 6. Summary ======
print(f"\n{'='*60}")
print(f"  v1.6.3 DOCX 同步完成")
print(f"{'='*60}")
print(f"  文件: {DOCX_ROOT}")
print(f"  大小: {final_size:,} bytes")
print(f"  边距: 2cm (四边)")
print(f"  版本: v1.6.3")
print(f"  备份: {backup_path}")
print(f"  三件套: .md✅ v1.6.3 / .json✅ v1.6.3 / .docx✅ v1.6.3")
