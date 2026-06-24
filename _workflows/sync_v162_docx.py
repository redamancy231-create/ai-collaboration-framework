#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Update AI协作项目全生命周期框架.docx to v1.6.2 — metadata + status footer + version header.

Targeted update — does NOT regenerate full content.
MD remains the authoritative version.
Run: PYTHONIOENCODING=utf-8 python _workflows/sync_v162_docx.py
"""

import os
import shutil
from datetime import datetime
from docx import Document

DOCX_PATH = "./AI协作项目全生命周期框架.docx"
BACKUP_PATH = "./AI协作项目全生命周期框架.docx.v161.backup"

def main():
    # Backup
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup: {BACKUP_PATH}")

    doc = Document(DOCX_PATH)

    # === 1. Core properties ===
    cp = doc.core_properties
    cp.version = "v1.6.2"
    cp.revision = 18
    cp.modified = datetime(2026, 6, 21)

    # === 2. Update version header in early paragraphs ===
    updated_header = False
    for i, para in enumerate(doc.paragraphs):
        if i > 15:
            break
        text = para.text
        if 'v1.6' in text and ('版本' in text or 'AI协作项目全生命周期框架' in text):
            for run in para.runs:
                if 'v1.6.1' in run.text:
                    run.text = run.text.replace('v1.6.1', 'v1.6.2')
                    updated_header = True
                    print(f"  Updated v1.6.1->v1.6.2 in paragraph {i}: '{text[:80]}...'")
                    break
                elif 'v1.6' in run.text and 'v1.6.1' not in run.text and 'v1.6.2' not in run.text:
                    run.text = run.text.replace('v1.6', 'v1.6.2')
                    updated_header = True
                    print(f"  Updated v1.6->v1.6.2 in paragraph {i}: '{text[:80]}...'")
                    break
    if not updated_header:
        print("  NOTE: Could not find version header paragraph to update")

    # === 3. Update status footer ===
    status_updated = False
    for para in reversed(doc.paragraphs):
        if '本文档状态' in para.text:
            new_status = (
                "本文档状态: v1.6.2 草案 (2026-06-21)。"
                "v1.6.2新增: §7.7被动观测——意外发现的发现机制(定义与概念边界+三次经验种子+8渠道扩展分类框架待实证"
                "+Failure Space 10种失效模式+硬约束+深度版Retrospect模板增强3可选字段)"
                "+ §9.11跨层可观测性设计(L0-L5各层关切+3原则)。"
                "经Codex GPT-5.5魔鬼代言人审查(写前,有条件支持)+Qwen qwen3.7-max完备性检查(写后,1处轻微错误已修正)。"
                "首次满足§2.6 Minor升级审查门(≥2后端)。"
                "v1.6.1: A2 Qwen qwen3.7-max跨模型复现写回(首次跨模型方向一致复现,证据二维M0→M2)。"
                "v1.6: §9.6.1二维证据等级+§9.10 MF三层模板+§4.1.1.1对照实验设计强制检查清单"
                "+§2.6维护流程+§1.8诚实声明+§9.9路径D+附录H反向交叉引用。"
                "⚠️ 本文档从.md源码转换，v1.6.2新增内容尚未完整渲染。以.md为权威版本。"
                "三件套: .md✅ v1.6.2 / .json✅ v1.6.2 / .docx⚠️ 元数据已更新, 全文待重新生成。"
            )
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_status
            else:
                para.add_run(new_status)
            status_updated = True
            print(f"Updated status footer paragraph")
            break

    if not status_updated:
        last_para = doc.paragraphs[-1]
        if 'v1.6' in last_para.text or '三件套' in last_para.text or '本文档状态' in last_para.text:
            new_status = (
                "本文档状态: v1.6.2 草案 (2026-06-21)。"
                "v1.6.2新增: §7.7被动观测+§9.11跨层可观测性设计。"
                "经Codex GPT-5.5魔鬼代言人+Qwen完备性检查两后端审查。"
                "v1.6.1: A2 Qwen跨模型复现写回。v1.6: 证据体系升级+维护性增强。"
                "⚠️ 以.md为权威版本。三件套: .md✅ v1.6.2 / .json✅ v1.6.2 / .docx⚠️ 元数据已更新。"
            )
            for run in last_para.runs:
                run.text = ""
            if last_para.runs:
                last_para.runs[0].text = new_status
            else:
                last_para.add_run(new_status)
            print(f"Updated last paragraph (fallback)")
            status_updated = True

    # === 4. Save ===
    doc.save(DOCX_PATH)
    print(f"DOCX synced to v1.6.2: {DOCX_PATH} ({os.path.getsize(DOCX_PATH)} bytes)")
    print("Note: v1.6.2 content is NOT fully rendered. Full regeneration requires pandoc pipeline.")
    print("三件套状态: .md✅ v1.6.2 / .json✅ v1.6.2 / .docx⚠️ 元数据已更新")

if __name__ == '__main__':
    main()
