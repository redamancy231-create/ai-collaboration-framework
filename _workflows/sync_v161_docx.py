#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Update AI协作项目全生命周期框架.docx to v1.6.1 — metadata + status footer + version header.

Targeted update — does NOT regenerate full content.
MD remains the authoritative version.
Run: PYTHONIOENCODING=utf-8 python _workflows/sync_v161_docx.py
"""

import os
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt

DOCX_PATH = "./AI协作项目全生命周期框架.docx"
BACKUP_PATH = "./AI协作项目全生命周期框架.docx.v16.backup"

def main():
    # Backup
    shutil.copy2(DOCX_PATH, BACKUP_PATH)
    print(f"Backup: {BACKUP_PATH}")

    doc = Document(DOCX_PATH)

    # === 1. Core properties ===
    cp = doc.core_properties
    cp.version = "v1.6.1"
    cp.revision = 17
    cp.modified = datetime(2026, 6, 20)

    # === 2. Update version header in early paragraphs ===
    # Search first 10 paragraphs for "v1.6" version header pattern
    updated_header = False
    for i, para in enumerate(doc.paragraphs):
        if i > 15:
            break
        text = para.text
        if 'v1.6' in text and ('版本' in text or 'AI协作项目全生命周期框架' in text):
            # Replace v1.6 → v1.6.1 in this paragraph's runs
            for run in para.runs:
                if 'v1.6' in run.text and 'v1.6.1' not in run.text:
                    run.text = run.text.replace('v1.6', 'v1.6.1')
                    updated_header = True
                    print(f"  Updated version in paragraph {i}: '{text[:60]}...'")
                    break
    if not updated_header:
        print("  NOTE: Could not find version header paragraph to update")

    # === 3. Update status footer (last paragraph) ===
    # Find the status paragraph — it's usually the last substantial paragraph
    # or the one containing "本文档状态"
    status_updated = False
    for para in reversed(doc.paragraphs):
        if '本文档状态' in para.text:
            new_status = (
                "本文档状态: v1.6.1 草案 (2026-06-20)。"
                "v1.6.1新增: A2 Qwen qwen3.7-max跨模型复现写回——首次跨模型方向一致复现, "
                "证据二维M0→M2, §4.1.1/§1.8/§6.3.1/§9.6.1更新。"
                "v1.6 P0新增: §9.6.1二维证据等级 + §9.10 MF三层模板 + §4.1.1.1对照实验设计强制检查清单。"
                "v1.6 P1新增: §2.6维护流程 + §1.8诚实声明 + §9.9路径D + 附录H反向交叉引用。"
                "经Codex GPT-5.5 v1.6初审→重审 + v1.6.1交叉验证(FAIL_WITH_CAVEATS→修正)。"
                "⚠️ 本文档从.md源码转换，v1.6.1新增内容尚未完整渲染。以.md为权威版本。"
                "三件套: .md✅ v1.6.1 / .json✅ v1.6.1 / .docx⚠️ 元数据已更新, 全文待重新生成。"
            )
            # Replace all runs in this paragraph
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_status
            else:
                para.add_run(new_status)
            status_updated = True
            print(f"Updated status footer paragraph")
            break

    if not status_updated:
        # Fallback: update last paragraph if it looks like a status block
        last_para = doc.paragraphs[-1]
        old_text = last_para.text[:80] if last_para.text else "(empty)"
        if 'v1.6' in last_para.text or '三件套' in last_para.text or '本文档状态' in last_para.text:
            new_status = (
                "本文档状态: v1.6.1 草案 (2026-06-20)。"
                "v1.6.1新增: A2 Qwen qwen3.7-max跨模型复现写回——首次跨模型方向一致复现, 证据二维M0→M2。"
                "v1.6新增: §9.6.1二维证据等级 + §9.10 MF三层模板 + §4.1.1.1对照实验设计强制检查清单 + "
                "§2.6维护流程 + §1.8诚实声明 + §9.9路径D + 附录H反向交叉引用。"
                "经Codex GPT-5.5初审→重审→v1.6.1交叉验证。"
                "⚠️ 以.md为权威版本。三件套: .md✅ v1.6.1 / .json✅ v1.6.1 / .docx⚠️ 元数据已更新。"
            )
            for run in last_para.runs:
                run.text = ""
            if last_para.runs:
                last_para.runs[0].text = new_status
            else:
                last_para.add_run(new_status)
            print(f"Updated last paragraph (fallback): '{old_text}...'")
            status_updated = True

    # === 4. Save ===
    doc.save(DOCX_PATH)
    print(f"DOCX synced to v1.6.1: {DOCX_PATH} ({os.path.getsize(DOCX_PATH)} bytes)")
    print("Note: full content regeneration requires md_to_docx_framework.py or pandoc.")

if __name__ == '__main__':
    main()
